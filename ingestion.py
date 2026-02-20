import os
import fitz  
import pdfplumber
from docx import Document
from typing import List, Optional
from langchain_text_splitters import RecursiveCharacterTextSplitter
from langchain_core.documents import Document as LCDocument
from langchain_huggingface import HuggingFaceEmbeddings
from langchain_chroma import Chroma
from dotenv import load_dotenv

load_dotenv()

def extract_text_from_pdf(file_path: str) -> str:
    """Extract text from PDF using pdfplumber (better for tables/columns)."""
    text = ""
    try:
        with pdfplumber.open(file_path) as pdf:
            for page in pdf.pages:
                page_text = page.extract_text()
                if page_text:
                    text += page_text + "\n\n"
    except Exception:
        doc = fitz.open(file_path)
        for page in doc:
            text += page.get_text() + "\n\n"
        doc.close()
    return text.strip()

def extract_text_from_docx(file_path: str) -> str:
    """Extract text from DOCX file."""
    doc = Document(file_path)
    paragraphs = [para.text for para in doc.paragraphs if para.text.strip()]
    return "\n\n".join(paragraphs)

def extract_text(file_path: str) -> str:
    """Auto-detect file type and extract text."""
    ext = os.path.splitext(file_path)[1].lower()
    if ext == ".pdf":
        return extract_text_from_pdf(file_path)
    elif ext in [".docx", ".doc"]:
        return extract_text_from_docx(file_path)
    else:
        raise ValueError(f"Unsupported file format: {ext}")

def chunk_text(
    text: str,
    file_name: str,
    chunk_size: int = None,
    chunk_overlap: int = None,
) -> List[LCDocument]:
    """Split text into overlapping chunks with metadata."""
    chunk_size = chunk_size or int(os.getenv("CHUNK_SIZE", 1000))
    chunk_overlap = chunk_overlap or int(os.getenv("CHUNK_OVERLAP", 200))
    splitter = RecursiveCharacterTextSplitter(
        chunk_size=chunk_size,
        chunk_overlap=chunk_overlap,
        separators=["\n\n", "\n", ". ", " ", ""],
    )
    chunks = splitter.split_text(text)
    documents = []
    for i, chunk in enumerate(chunks):
        documents.append(
            LCDocument(
                page_content=chunk,
                metadata={
                    "source": file_name,
                    "chunk_index": i,
                    "total_chunks": len(chunks),
                },
            )
        )
    return documents

def get_embeddings():
    """Initialize HuggingFace embeddings (مجاني - بيشتغل محلياً)."""
    return HuggingFaceEmbeddings(                         
        model_name="sentence-transformers/all-MiniLM-L6-v2",
        model_kwargs={"device": "cpu"},
        encode_kwargs={"normalize_embeddings": True},
    )

def ingest_document(
    file_path: str,
    collection_name: str = "contracts",
    persist_directory: str = None,
) -> Chroma:
    """
    Full ingestion pipeline:
    1. Extract text
    2. Chunk text
    3. Create embeddings
    4. Store in Chroma vector DB
    """
    persist_directory = persist_directory or os.getenv("CHROMA_DB_PATH", "./chroma_db")
    file_name = os.path.basename(file_path)
    print(f"Extracting text from: {file_name}")
    text = extract_text(file_path)
    if not text:
        raise ValueError("No text could be extracted from the document.")
    print(f"Chunking text ({len(text)} chars)...")
    documents = chunk_text(text, file_name)
    print(f"Created {len(documents)} chunks")
    print("Creating embeddings and storing in Chroma...")
    embeddings = get_embeddings()
    vectorstore = Chroma.from_documents(
        documents=documents,
        embedding=embeddings,
        collection_name=collection_name,
        persist_directory=persist_directory,
    )
    print(f"Ingestion complete! {len(documents)} chunks stored.")
    return vectorstore

def load_vectorstore(
    collection_name: str = "contracts",
    persist_directory: str = None,
) -> Chroma:
    """Load existing Chroma vector store."""
    persist_directory = persist_directory or os.getenv("CHROMA_DB_PATH", "./chroma_db")
    embeddings = get_embeddings()                          
    return Chroma(
        collection_name=collection_name,
        embedding_function=embeddings,
        persist_directory=persist_directory,
    )